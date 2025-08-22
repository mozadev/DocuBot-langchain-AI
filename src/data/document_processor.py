"""
Procesador de documentos para DocuBot AI.
Maneja PDF, DOCX y otros formatos de documentos.
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from src.core.logger import logger, log_function_call
from config.settings import settings

class DocumentProcessor:
    """
    Procesador de documentos que extrae texto y lo divide en chunks.
    Implementa las mejores prácticas de procesamiento de documentos.
    """
    
    def __init__(self):
        """Inicializa el procesador de documentos."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        logger.info("DocumentProcessor inicializado")
    
    @log_function_call
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extrae texto de un archivo PDF.
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            Texto extraído del PDF
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Texto extraído de PDF: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de PDF {file_path}: {str(e)}")
            raise
    
    @log_function_call
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extrae texto de un archivo DOCX.
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            Texto extraído del DOCX
        """
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            logger.info(f"Texto extraído de DOCX: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de DOCX {file_path}: {str(e)}")
            raise
    
    @log_function_call
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extrae texto de un archivo basándose en su extensión.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Texto extraído del archivo
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif file_extension == '.docx':
            return self.extract_text_from_docx(str(file_path))
        else:
            raise ValueError(f"Formato de archivo no soportado: {file_extension}")
    
    @log_function_call
    def split_text_into_chunks(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[LangChainDocument]:
        """
        Divide el texto en chunks para procesamiento posterior.
        
        Args:
            text: Texto a dividir
            metadata: Metadatos adicionales para los chunks
            
        Returns:
            Lista de documentos LangChain
        """
        try:
            # Crear documento LangChain
            doc = LangChainDocument(
                page_content=text,
                metadata=metadata or {}
            )
            
            # Dividir en chunks
            chunks = self.text_splitter.split_documents([doc])
            
            logger.info(f"Texto dividido en {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error dividiendo texto en chunks: {str(e)}")
            raise
    
    @log_function_call
    def process_document(self, file_path: str) -> List[LangChainDocument]:
        """
        Procesa un documento completo: extrae texto y lo divide en chunks.
        
        Args:
            file_path: Ruta al archivo a procesar
            
        Returns:
            Lista de chunks del documento
        """
        try:
            # Extraer texto
            text = self.extract_text_from_file(file_path)
            
            # Crear metadatos
            metadata = {
                'source': file_path,
                'filename': Path(file_path).name,
                'file_type': Path(file_path).suffix,
                'file_size': os.path.getsize(file_path)
            }
            
            # Dividir en chunks
            chunks = self.split_text_into_chunks(text, metadata)
            
            logger.info(f"Documento procesado exitosamente: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error procesando documento {file_path}: {str(e)}")
            raise
